import magic # For filetype detection
from pathlib import Path # For filetype detection
import PyPDF2 # For scanning PDF
import PyPDF2.errors # For scanning PDF
import pdfplumber # For scanning PDF
import re # For finding specific strings in files
from docx import Document # For DOCX processing
from pptx import Presentation # For PPTX processing
from langdetect import detect # Detecting language in DOCX
from oletools.olevba import VBA_Parser # Detecting Macros in DOCX
import zipfile # Compressed file detection
import py7zr # Compressed file detection
import subprocess # For Strings command
import pefile # For PE header analysis
import peutils # For PE packer analysis
import os # For PE header analysis
import datetime # For PE header analysis
import webbrowser # For displaying results as HTML
import threading # For multithreading
from threading import Lock # For fluent HTML display

# --------------------------------------HTML--------------------------------------

def write_html(html_content, filename):
    with open(filename, 'w') as file:
        file.write(html_content)
    webbrowser.open('file://' + os.path.realpath(filename))

# ---------------------------------------.PDF--------------------------------------- 

def is_password_protected_pdf(filepath): # Check if PDF is encyrpted.
    try:
        with open(filepath, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            if pdf_reader.is_encrypted:
                return True
            else:
                return False
    except Exception:
        return True

def get_pdf_contents(filepath,html_content): # Read contents of the PDF file.
    urls = []
    ips = []
    domains = []
    url_pattern = re.compile(r'(https?:\/\/(?:www\.|(?!www))[a-zA-Z0-9][a-zA-Z0-9-]+[a-zA-Z0-9]\.[^\s]{2,}|www\.[a-zA-Z0-9][a-zA-Z0-9-]+[a-zA-Z0-9]\.[^\s]{2,}|https?:\/\/(?:www\.|(?!www))[a-zA-Z0-9]+\.[^\s]{2,}|www\.[a-zA-Z0-9]+\.[^\s]{2,})')
    ip_pattern = re.compile(r'^((25[0-5]|(2[0-4]|1\d|[1-9]|)\d)\.){3}(25[0-5]|(2[0-4]|1\d|[1-9]|)\d)$')
    domain_pattern = re.compile(r'^((?!-)[A-Za-z0–9-]{1, 63}(?<!-)\.)+[A-Za-z]{2, 6}$')
    with pdfplumber.open(filepath) as pdf:
        pages = pdf.pages
        for page in pages:
            text = page.extract_text()
            if text:
                urls.extend(url_pattern.findall(text))
                ips.extend(ip_pattern.findall(text))
                domains.extend(domain_pattern.findall(text))
    html_content = html_content + """<h1>Strings</h1><table>
    <tr><th>Category</th><th>String</th></tr>"""
    for i in urls:
        html_content = html_content + f"""<tr><td>URL</td><td>{i}</td></tr>"""
    for i in ips:
        html_content = html_content + f"""<tr><td>IP</td><td>{i}</td></tr>"""
    for i in domains:
        html_content = html_content + f"""<tr><td>Domain</td><td>{i}</td></tr>"""
    html_content = html_content + """</table>"""
    return html_content

# ------------------------------------COMPRESSED------------------------------------

def check_zip_password(file_path,html_content): # ZIP password
    try:
        with zipfile.ZipFile(file_path) as zf:
            if zf.testzip() is None:
                return html_content + """<h1>General Information</h1><h2>This ZIP file is not password protected.</h2>"""
    except Exception:
        return html_content + """<h1>General Information</h1><h2>This ZIP file is password protected.</h2>"""

def check_7z_password(file_path,html_content):  # 7ZIP password
    try:
        with py7zr.SevenZipFile(file_path, mode='r') as z:
            z.extractall()
            return html_content + """<h1>General Information</h1><h2>This 7ZIP file is not password protected.</h2>"""
    except Exception:
        return html_content + """<h1>General Information</h1><h2>This 7ZIP file is password protected.</h2>"""

# ---------------------------------------DOCX---------------------------------------

def is_encrypted_docx(filepath): # Check for encyrption
    try:
        Document(filepath)
        return False
    except Exception:
        return True

def is_encrypted_pptx(filepath): # Check for encyrption
    try:
        Presentation(filepath)
        return False
    except Exception:
        return True

def is_password_protected_docx(filepath): # Check for password
    try:
        doc = Document(filepath)
        doc.paragraphs
        return False
    except Exception:
        return True

def detect_language(filepath,html_content): # Detect language
    doc = Document(filepath)
    context = []
    for i in doc.paragraphs:
        context.append(i.text)
    text = ' '.join(context)
    html_content = html_content + f"""<h1>General Information</h1><h2>Document language: {detect(text)}</h2>"""
    return html_content

def has_macros(filepath,html_content):
    vbaparser = VBA_Parser(filepath)
    if(vbaparser.detect_vba_macros()):
        html_content = html_content + """<h2>Document contains macros which may be harmful or malicious!</h2>"""
    else:
        html_content = html_content + """<h2>Document does not contain any macros.</h2>"""
    return html_content
# -------------------------------------FILETYPE-------------------------------------

def get_extension_name(filepath): # Get filetype based on filename
    return Path(filepath).suffix[1:]

def get_extension_magic(filepath): # Get filetype based on magic value
    magic_obj = magic.Magic(mime=True)
    file_type = magic_obj.from_file(filepath)
    file_type = file_type.split('/')[1]
    return file_type

def get_filetype(filename,html_content): # Main function to handle filetype operations
    try:
        file_extension_name = get_extension_name(filename)
        file_extension_magic = get_extension_magic(filename)
        html_content = html_content + """<h1>Filetype Information</h1>"""
        html_content = html_content + f"""<h2>This file appears to have the extension: {file_extension_name}</h2>"""
        html_content = html_content + f"""<h2>The file has the extension: {file_extension_magic}</h2>"""
    except Exception as err:
        print(err)
    return html_content,file_extension_magic

# ----------------------------------------PE----------------------------------------

def extract_strings(filepath,html_content):
    try:
        # Run the strings command
        result = subprocess.run(['strings', filepath], capture_output=True, text=True)
        # Return the output if the command was successful
        if result.returncode == 0:
            urls = []
            ips = []
            domains = []
            url_pattern = re.compile(r'(https?:\/\/(?:www\.|(?!www))[a-zA-Z0-9][a-zA-Z0-9-]+[a-zA-Z0-9]\.[^\s]{2,}|www\.[a-zA-Z0-9][a-zA-Z0-9-]+[a-zA-Z0-9]\.[^\s]{2,}|https?:\/\/(?:www\.|(?!www))[a-zA-Z0-9]+\.[^\s]{2,}|www\.[a-zA-Z0-9]+\.[^\s]{2,})')
            ip_pattern = re.compile(r'^((25[0-5]|(2[0-4]|1\d|[1-9]|)\d)\.){3}(25[0-5]|(2[0-4]|1\d|[1-9]|)\d)$')
            domain_pattern = re.compile(r'^((?!-)[A-Za-z0–9-]{1, 63}(?<!-)\.)+[A-Za-z]{2, 6}$')
            stdout_string = result.stdout.split()
            for i in stdout_string:
                urls.extend(url_pattern.findall(i))
                ips.extend(ip_pattern.findall(i))
                domains.extend(domain_pattern.findall(i))
            html_content = html_content + """<h1>Strings</h1><table>
            <tr><th>Category</th><th>String</th></tr>"""
            for i in urls:
                html_content = html_content + f"""<tr><td>URL</td><td>{i}</td></tr>"""
            for i in ips:
                html_content = html_content + f"""<tr><td>IP</td><td>{i}</td></tr>"""
            for i in domains:
                html_content = html_content + f"""<tr><td>Domain</td><td>{i}</td></tr>"""
            html_content = html_content + """</table>"""
            return html_content
        else:
            html_content = html_content + "<h1>Strings</h1><h2>Strings command yielded no results.</h2>"
            return html_content
    except:
        return html_content

def analyze_header(filepath,html_content):
    pe = pefile.PE(filepath)
    data = {
        'Architecture': 'x86-x64' if pe.FILE_HEADER.Machine == 0x8664 else 'x86',
        'Entropy': 0, #Value adjusted later
        'File Size': os.path.getsize(filepath),
        'Number of Sections': len(pe.sections),
        'Compilation Date': datetime.datetime.fromtimestamp(pe.FILE_HEADER.TimeDateStamp).strftime('%Y-%m-%d %H:%M:%S'),
    }

    sections = []
    for section in pe.sections:
        section_data = {
            'Name': section.Name.decode().strip('\x00'),
            'Entropy': section.get_entropy(),
            'Virtual Size': section.Misc_VirtualSize,
            'Raw Size': section.SizeOfRawData
        }
        sections.append(section_data)

    data['Entropy'] = calculate_overall_entropy(sections)

    imported_dlls = []
    if hasattr(pe, 'DIRECTORY_ENTRY_IMPORT'):
        for entry in pe.DIRECTORY_ENTRY_IMPORT:
            dll_name = entry.dll.decode()
            imported_dlls.append(dll_name)
    html_addition = f"""    
        <h1>General Information</h1>
            <h2>Architecture: {data['Architecture']}</h2>
            <h2>Entropy: {data['Entropy']}</h2>
            <h2>File Size (In Bytes): {data['File Size']}</h2>
            <h2>Sections: {data['Number of Sections']}</h2>
            <h2>Compilation Date: {data['Compilation Date']}</h2>
        <h1>Imported DLLs</h1>
        <table>
            <tr><th>DLL</th></tr>
    """
    html_content = html_content + html_addition
    for dll in imported_dlls:
        html_content = html_content + f"<tr><td>{dll}</td>"
    html_content = html_content + "</table><h1>Section Information</h1><table><tr><th>Name</th><th>Entropy</th><th>Virtual Size</th><th>Raw Size</th></tr>"
    for section in sections:
        html_content = html_content + f"<tr><td>{section['Name']}</td><td>{section['Entropy']}</td><td>{section['Virtual Size']}</td><td>{section['Raw Size']}</td></tr>"
    html_content = html_content + "</table>"
    return html_content,sections

def calculate_overall_entropy(sections):
    entropy = 0.0
    val = 0
    for i in sections:
        if (entropy == 0.0):
            entropy = i['Entropy']
            val = i['Raw Size']
        else:
            entropy = (entropy*val + i['Entropy']*i['Raw Size']) / (val + i['Raw Size'])
            val = val + i['Raw Size']
    return entropy

def analyze_packing(sections,html_content): # Check for packing status
    packing_probability = 0
    for i in sections:
        if(i['Entropy'] > 6.0):
            html_content = html_content + f"""<h2>- The section {i['Name']} has high entropy value of {i['Entropy']}</h2>"""
            if(i['Entropy'] > 7.5):
                packing_probability = 100
            elif(i['Entropy'] > 7.0):
                packing_probability = 90
            elif(i['Entropy'] > 6.5):
                packing_probability = 75
            else:
                packing_probability = 50
        if(i['Raw Size'] == 0): # Avoiding division by Zero
            i['Raw Size'] = 1
        if(i['Virtual Size'] / i['Raw Size'] > 10.0):
            html_content = html_content + f"""<h2>- The section {i['Name']} has high Virtual to Raw size ratio with value {i['Virtual Size'] / i['Raw Size']}.</h2>"""
    if packing_probability > 70:
        html_content = html_content + "<h2>This file is packed using an algortihm that is not in the signature database.</h2>"
    elif 75 >= packing_probability >= 50:
        html_content = html_content + "<h2>This file is likely packed using an algortihm that is not in the signature database.</h2>"
    else:
        html_content = html_content + "<h2>This file is not packed.</h2>"
    return html_content

def load_signature_database(filepath): # Helper function for find_packer()
    with open(filepath, 'r', encoding='latin-1') as file:
        signature_data = file.read()
    return peutils.SignatureDatabase(data=signature_data)

def find_packer(filepath,sections,html_content):
    try:
        pe = pefile.PE(filepath)
        signatures = load_signature_database("UserDB.TXT") # signature database from http://woodmann.com/BobSoft/Files/Other/UserDB.zip
        result = signatures.match(pe, ep_only=True)
        if result:
            html_content = html_content + f"""<h1>Packer Information</h1><h2>This file was packed using the following packer: {result}</h2>"""
            return html_content
        else:
            html_content = html_content + f"""<h1>Packer Information</h1><h2>No packing algorithm matches the file, manual analysis found the following result:<h2>"""
            html_content = analyze_packing(sections,html_content)
            return html_content
    except Exception as err:
        print(err)
        return html_content

# ---------------------------------------MAIN---------------------------------------

def main_func(filename):
    html_content = f"""
        <html>
        <head>
            <title>{filename}</title>
            """+"""<style>
                body { font-family: Arial, sans-serif; }
                table { width: 100%; border-collapse: collapse; }
                th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
            </style>
        </head>
        <body>"""
    html_content,filetype = get_filetype(filename,html_content)
    if (filetype == "pdf"):
        if( not is_password_protected_pdf(filename)):
            html_content = get_pdf_contents(filename,html_content)
        else:
            html_content = html_content + """<h1>General Information</h1><h2>PDF File is password protected.</h2>"""
    elif (filetype == "vnd.openxmlformats-officedocument.wordprocessingml.document"):
        if(not is_encrypted_docx(filename)):
            if(not is_password_protected_docx(filename)):
                html_content = detect_language(filename,html_content)
                html_content = has_macros(filename,html_content)
        else:
            html_content = html_content + """<h1>General Information</h1><h2>DOCX File is password protected.</h2>"""
    elif (filetype == "vnd.openxmlformats-officedocument.presentationml.presentation"):
        if(not is_encrypted_pptx(filename)):
            html_content = has_macros(filename,html_content)      
        else:
            html_content = html_content + """<h1>General Information</h1><h2>PPTX File is password protected.</h2>"""          
    elif(filetype == "zip"):
        html_content = check_zip_password(filename,html_content)
    elif(filetype == "x-7z-compressed"):
        html_content = check_7z_password(filename,html_content)
    elif(filetype == "executable" or filetype == "x-dosexec"):
        html_content, sections = analyze_header(filename,html_content)
        html_content = find_packer(filename,sections,html_content)
        html_content = extract_strings(filename,html_content)
    html_content = html_content + """</body></html>"""
    mutex_lock.acquire()
    write_html(html_content,'./outputs/'+filename+'.html')
    mutex_lock.release()

print("Enter filenames that you want to analyze, for entry of multiple files separate them by ',' character.")
filenames = input("Enter Filename(s): ")
threads = []
mutex_lock = Lock()
if (filenames.count(",") > 0):
    for i in filenames.split(","):
        i = i.strip()
        thread = threading.Thread(target=main_func, args=(i,))
        threads.append(thread)
        thread.start()
else:
    main_func(filenames.strip())

for thread in threads:
    thread.join()